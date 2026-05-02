"""
Module chứa các hàm đọc file từ định dạng .docx và .txt.
"""
import os
from typing import Optional
from docx import Document

def load_docx(file_path: str) -> Optional[str]:
    """
    Đọc file .docx và trả về toàn bộ text.
    """
    try:
        doc = Document(file_path)
        full_text = []
        # Đọc các đoạn văn
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        # Đọc các bảng (nếu có)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text.strip())
        return "\n".join(full_text)
    except Exception as e:
        print(f"Lỗi đọc file {file_path}: {e}")
        return None

def load_text(file_path: str) -> Optional[str]:
    """
    Đọc file .txt thông thường.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Lỗi đọc file {file_path}: {e}")
        return None

def load_file(file_path: str) -> Optional[str]:
    """
    Tự động phát hiện định dạng và gọi loader tương ứng.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.docx':
        return load_docx(file_path)
    elif ext == '.txt':
        return load_text(file_path)
    else:
        print(f"Định dạng không hỗ trợ: {file_path}")
        return None